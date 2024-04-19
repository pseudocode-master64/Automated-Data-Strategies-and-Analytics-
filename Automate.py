import pandas as pd
import csv
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING



# Turn to CSV File
inputExcelFile ="Liberal_Arts.xlsx"

excelFile = pd.read_excel (inputExcelFile)

excelFile.to_csv ("ResultCsvFile.csv", index = None, header=True)

dataframeObject = pd.DataFrame(pd.read_csv("ResultCsvFile.csv"))

input_file = "ResultCsvFile.csv"

# Create Word doc Section
def createSection(C_ID, C_Name, C_Prof, C_Type, C_Time, C_Grade, C_OH, C_Speak, C_HWA, C_R, C_HWD, C_TD, C_LR, C_ATT, C_DS, C_EC, C_ProfQua, C_C):

    # Section 1
    p = doc.add_paragraph()

    C_ID = C_ID.upper()
    C_Name = C_Name.upper()

    p.add_run(f'{C_ID}: {C_Name}').bold = True

    # Section 2

    C_Prof = C_Prof.split(" ")
    if len(C_Prof ) > 1:
        C_Prof = str(C_Prof[0].capitalize()+ " " + C_Prof[1].capitalize())
    else:
        C_Prof = str(C_Prof[0].capitalize())

    if C_Type == "In person":
        C_Type =  "线下"
    elif C_Type == "Online":
        C_Type = "线上"
    else:
        C_Type = "线上+线下"

    C_Time = C_Time.split(" ")
    C_Time = str(C_Time[1] + " " + C_Time[0])


    if C_Grade == "":
        C_Grade = "N/A"

    p.add_run(f'''
任课教授：{C_Prof}
授课形式：{C_Type}
上该课的学期：{C_Time}
''')

    p.add_run(f'Expected Letter Grade: {C_Grade}').bold = True

    # Section 3

    p.add_run(f'''
评分：''')

    p = doc.add_paragraph()

    p_OH = doc.add_paragraph(style='List Bullet')
    p_OH.add_run(f'教授Office Hour：{C_OH}.0')

    p_Speak = doc.add_paragraph(style='List Bullet')
    p_Speak.add_run(f'教授口齿清晰程度：{C_Speak}.0')

    p_HWA = doc.add_paragraph(style='List Bullet')
    p_HWA.add_run(f'课程作业量：{C_HWA}.0')

    p_R = doc.add_paragraph(style='List Bullet')
    p_R.add_run(f'课程阅读量：{C_R}.0')

    p_HWD = doc.add_paragraph(style='List Bullet')
    p_HWD.add_run(f'作业难度：{C_HWD}.0')

    p_TD = doc.add_paragraph(style='List Bullet')
    p_TD.add_run(f'考试难度：{C_TD}.0')

    if C_LR == "是":
        C_LR = "有"
    else:
        C_LR = "无"

    if C_ATT == "是":
        C_ATT = "记录"
    else:
        C_ATT = "不记录"

    if C_DS == "是":
        C_DS = "有"
    else:
        C_DS = "无"

    if C_EC == "是":
        C_EC = "有"
    else:
        C_EC = "无"

    p_LR = doc.add_paragraph(style='List Bullet')
    p_LR.add_run(f'Lecture Recording：{C_LR}')

    p_ATT = doc.add_paragraph(style='List Bullet')
    p_ATT.add_run(f'Lecture Attendance：{C_ATT}')

    p_DS = doc.add_paragraph(style='List Bullet')
    p_DS.add_run(f'Discussion Session：{C_DS}')

    p_EC = doc.add_paragraph(style='List Bullet')
    p_EC.add_run(f'Extra Credit：{C_EC}')

    p_ProfQua = doc.add_paragraph(style='List Bullet')
    p_ProfQua.add_run(f'教授授课水平：{C_ProfQua}.0')


    # Section 4
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)

    p.add_run(f'''详细评价/建议：''')
    p = doc.add_paragraph()
    p_C = doc.add_paragraph(style='List Bullet')
    p_C.add_run(f'{C_C} \n')


    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(12)
        # Set the line spacing to single spacing
# Read csv file and generate word doc
with open (input_file, 'r') as infile :
    reader = csv.reader(infile)
    next(reader)
    doc = Document()
    for row in reader: # Read each row one by one
        C_ID, C_Name, C_Prof, C_Type, C_Time, C_Grade, C_OH, C_Speak, C_HWA, C_R, C_HWD, C_TD, C_LR, C_ATT, C_DS, C_EC, C_ProfQua, C_C = row[0],row[1], row[2], row[3], row[4], row[5], row[6], row[7], row[8], row[9], row[10], row[11], row[12], row[13], row[14], row[15], row[16], row[17]
        createSection(C_ID, C_Name, C_Prof, C_Type, C_Time, C_Grade, C_OH, C_Speak, C_HWA, C_R, C_HWD, C_TD, C_LR, C_ATT, C_DS, C_EC, C_ProfQua, C_C)
    
    doc.save('Result.docx')

# Caution:
# Use word document to change all bullet point type
# Change title font size
# Delete space before 建议 section
# Need to standardize input